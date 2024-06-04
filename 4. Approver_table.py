import os
from docx import Document

def find_approver_table(doc, start_text, end_text):
    start_found = False
    end_found = False
    for para in doc.paragraphs:
        if not start_found and para.text.strip() == start_text:
            start_found = True
        elif start_found and para.text.strip() == end_text:
            end_found = True
            break
    if start_found and end_found:
        for table in doc.tables:
            if table.rows and table.rows[0].cells:
                first_row_cells = [cell.text.strip() for cell in table.rows[0].cells]
                if first_row_cells[0].lower() == "name" and first_row_cells[-1].lower() == "version":
                    return table
    return None

def validate_approver_table(approver_table):
    if approver_table:
        for row in approver_table.rows[1:]:
            name = row.cells[0].text.strip()
            title_department = row.cells[2].text.strip()
            date_of_review = row.cells[3].text.strip()
            if not name or not title_department or not date_of_review:
                print("Name, Title/Department, or Date of Review not filled in for an approver.")
                return False
        print("Name, Title/Department, and Date of Review filled in for all approvers.")
        return True
    else:
        print("Approver table not found.")
        return False

# Path to the directory containing the Word documents
doc_directory = r"C:\Users\lazare.kolebka\OneDrive - Accenture\Desktop\Yara"

# Parameters for finding the reviewer table
start_text = "Reviewers"
end_text = "This document has been reviewed by:"

# Iterate over all Word documents in the directory
for filename in os.listdir(doc_directory):
    if filename.endswith(".docx"):
        doc_path = os.path.join(doc_directory, filename)
        doc = Document(doc_path)
        
        approver_table = find_approver_table(doc, start_text, end_text)
        print(f"Validating approver table for document {filename}:")
        validation_status = validate_approver_table(approver_table)
        print(f"Finished processing document: {doc_path}\n")

print("Table extraction and validation completed.")
