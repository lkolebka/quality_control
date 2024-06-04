import os
from docx import Document
from collections import defaultdict

def read_wricef_tables(doc_path):
    print(f"Reading WRICEF tables from document: {doc_path}")
    doc = Document(doc_path)
    wricef_tables = []

    for table in doc.tables:
        if table.rows and len(table.rows) > 1 and table.rows[1].cells:
            # Check second row for WRICEF ID header
            second_row_cells = [cell.text.strip() for cell in table.rows[1].cells]
            print(f"Checking table with headers: {second_row_cells}")
            if second_row_cells[0].lower() == "wricef id":
                wricef_tables.append(table)

    return wricef_tables

def validate_wricef_table(wricef_table):
    correct_count = 0
    value_counts = defaultdict(int)
    # Start from the third row to skip the header and ensure it is valid
    for row in wricef_table.rows[2:]:
        wricef_id = row.cells[0].text.strip().lower()
        try:
            # Check if WRICEF ID is a valid integer
            int(wricef_id)
            correct_count += 1
        except ValueError:
            value_counts[wricef_id] += 1
    return correct_count, value_counts

def print_table(table):
    table_content = []
    for row in table.rows:
        row_content = [cell.text.strip() for cell in row.cells]
        table_content.append(" | ".join(row_content))
    return "\n".join(table_content)

# Path to the directory containing the Word documents
doc_directory = r"C:\Users\lazare.kolebka\OneDrive - Accenture\Desktop\Yara"

# Iterate over all Word documents in the directory
for filename in os.listdir(doc_directory):
    if filename.endswith(".docx"):
        doc_path = os.path.join(doc_directory, filename)
        doc = Document(doc_path)
        
        wricef_tables = read_wricef_tables(doc_path)
        
        if wricef_tables:
            total_correct = 0
            total_values = defaultdict(int)
            for table in wricef_tables:
                print(f"Table in document {filename}:\n{print_table(table)}\n")
                correct_count, value_counts = validate_wricef_table(table)
                total_correct += correct_count
                for key, count in value_counts.items():
                    total_values[key] += count
            
            print(f"Document {filename}:")
            print(f"{len(wricef_tables)} WRICEF table(s) found:")
            print(f"{total_correct} IDs are correct")
            for value, count in total_values.items():
                print(f"{count} IDs are '{value}'")
        else:
            print(f"Document {filename}: No WRICEF tables found.")
        print(f"Finished processing document: {doc_path}\n")

print("WRICEF table extraction and validation completed.")
