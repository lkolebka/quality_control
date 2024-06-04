import os
import pandas as pd
from docx import Document
from collections import defaultdict
from helper.wricef_validation_5 import read_wricef_tables, validate_wricef_table
from helper.approvers_validation_4 import find_approver_table, validate_approver_table
from helper.open_point_validation_6 import check_open_points

def check_blue_text(doc):
    blue_text_found = False
    blue_text_paragraphs = []
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.color and run.font.color.rgb == (0, 112, 192):  # RGB value
                blue_text_paragraphs.append(paragraph.text.strip())
                blue_text_found = True
    return blue_text_found, blue_text_paragraphs

def check_headers(doc):
    header_issues_found = False
    header_issue_details = []

    def get_next_valid_element(index):
        while index < len(doc.paragraphs):
            paragraph = doc.paragraphs[index]
            if paragraph.text.strip():
                return 'paragraph', paragraph.text.strip()
            index += 1
        # Check for tables
        for table in doc.tables:
            if doc.element.body.index(table._element) == index:
                return 'table', None
        return None, None

    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.style.name in ['Heading 1', 'Heading 2']:
            # Check the next valid element
            next_element_type, next_element_content = get_next_valid_element(i + 1)

            if next_element_type == 'paragraph' and (len(next_element_content) > 5 or next_element_content.lower() in ["n/a", "not applicable"]):
                valid_content = True
            elif next_element_type == 'table':
                valid_content = True
            else:
                valid_content = False

            if not valid_content:
                header_issues_found = True
                header_issue_details.append(f"{paragraph.style.name} '{paragraph.text}' has invalid or missing content after it.")

    return header_issues_found, header_issue_details

def extract_version_number(doc):
    version_numbers = []
    for section in doc.sections:
        header = section.header
        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "V." in cell.text:
                        version_numbers.append(cell.text)
    return version_numbers

def check_template_points(doc):
    template_points = [
        "Business Process Description",
        "Business Process Diagram",
        "Process Step Detailed Requirements & Solution",
        "Business Process Requirements",
        "Locations Where this Business Process is Performed",
        "Operational Decisions or Logic within the Process",
        "Legal Considerations and Company-Specific Policies",
        "Reference to Process KPIs, Measurements, etc.",
        "Process Interfaces",
        "System Integration Points",
        "Improvements Required to Satellite Systems and Applications",
        "Potential Future Process Improvements (out of scope for this implementation)",
        "Functional Solution Design",
        "Organization Structure Considerations",
        "Master Data Considerations (including all relevant data relationships)",
        "System Configuration Considerations",
        "Technical/Development Related Items",
        "Workflow",
        "Reporting (operational and analytical)",
        "Interfaces",
        "Data Conversion / Historical Data",
        "Enhancements",
        "Output (e.g. Forms)",
        "Service-Oriented Architecture/Composition",
        "Business Domain",
        "Process Component",
        "Service Model",
        "Service Operation Candidate",
        "Authorization",
        "Organizational Change Related Items",
        "Change Impact Summary",
        "Training Requirements",
        "Internal Control",
        "Interim Period Considerations",
        "Open Points",
        "Appendix",
        "Appendix A - Business Requirements"
    ]
    doc_points = [paragraph.text.strip() for paragraph in doc.paragraphs if paragraph.style.name.startswith('Heading')]
    missing_points = [point for point in template_points if point not in doc_points]
    return missing_points

def check_document_quality(doc_path):
    print(f".......................Checking document: {doc_path}...............")
    doc = Document(doc_path)
    report = {
        "document_title": os.path.basename(doc_path),
        "1. blue_text_check": "OK",
        "1. blue_text_details": "",
        "2. header_text_check": "OK",
        "2. header_text_details": "",
        "3. version_number": "Not Found",
        "4. approvers_table_check": "OK",
        "4. approvers_table_details": "",
        "5. ricef_check": "OK",
        "5. ricef_details": "",
        "6. point_9_check": "OK",
        "6. point_9_details": "",
        "7. template_points_check": "OK",
        "7. template_points_details": ""
    }

    # 1. Check for blue text
    blue_text_found, blue_text_paragraphs = check_blue_text(doc)
    if blue_text_found:
        report["1. blue_text_check"] = "Not OK"
        report["1. blue_text_details"] = "; ".join(blue_text_paragraphs)
        print(f"Check 1: ❌ Blue text found in document: {doc_path}")
    else:
        print(f"Check 1: ✅ No blue text detected in document: {doc_path}")

    # 2. Check headers
    header_issues_found, header_issue_details = check_headers(doc)
    if header_issues_found:
        report["2. header_text_check"] = "Not OK"
        report["2. header_text_details"] = "; ".join(header_issue_details)
    else:
        print(f"Check 2: ✅ All headers have valid text in document: {doc_path}")

    # 3. Extract version number from header
    version_numbers = extract_version_number(doc)
    if version_numbers:
        report["3. version_number"] = ", ".join(version_numbers)
        print(f"Check 3: ✅ Version number found in header of document: {doc_path} - Version: {', '.join(version_numbers)}")
    else:
        print(f"Check 3: ❌ Version number not found in header of document: {doc_path}")

    # 4. Check approvers table
    approver_table = find_approver_table(doc, "This document has been approved by:", "Distribution list")
    approver_table_valid = validate_approver_table(approver_table)
    if not approver_table_valid:
        report["4. approvers_table_check"] = "Not OK"
        report["4. approvers_table_details"] = "Name, Title/Department, or Date of Review not filled in for an approver."
        print(f"Check 4: ❌ Approvers table validation failed in document: {doc_path}")
    else:
        report["4. approvers_table_details"] = "All approvers have Name, Title/Department, and Date of Review filled in."
        print(f"Check 4: ✅ Approvers table validation passed in document: {doc_path}")

    # 5. Check WRICEF table
    wricef_tables = read_wricef_tables(doc)
    if wricef_tables:
        total_correct = 0
        total_values = defaultdict(int)
        for table in wricef_tables:
            correct_count, value_counts = validate_wricef_table(table)
            total_correct += correct_count
            for key, count in value_counts.items():
                total_values[key] += count
        total_rows = sum(len(table.rows) - 2 for table in wricef_tables)
        if total_correct == total_rows:
            print("Check 5: ✅ All WRICEF IDs are correct.")
            report["5. ricef_check"] = "OK"
        else:
            report["5. ricef_check"] = "Not OK"
            report["5. ricef_details"] = "; ".join(f"{count} IDs are '{value}'" for value, count in total_values.items())
            print(f"Check 5: ❌ Missing valid WRICEF IDs in document: {doc_path}")
    else:
        report["5. ricef_check"] = "Not OK"
        report["5. ricef_details"] = "No WRICEF tables found."
        print(f"Check 5: ❌ No WRICEF tables found in document: {doc_path}")

    # 6. Check point 9. open point
    if check_open_points(doc):
        report["6. point_9_check"] = "OK"
        report["6. point_9_details"] = "Link found in Open Points section."
    else:
        report["6. point_9_check"] = "Not OK"
        report["6. point_9_details"] = "No link found in Open Points section."

    # 7. Check all template points are present
    missing_points = check_template_points(doc)
    if missing_points:
        report["7. template_points_check"] = "Not OK"
        report["7. template_points_details"] = "Missing " + ", ".join(missing_points)
        print(f"Check 7: ❌ Missing template points in document: {doc_path} - Missing points: {', '.join(missing_points)}")
    else:
        print(f"Check 7: ✅ All template points are present in document: {doc_path}")

    return report

def main():
    # Path to the directory containing the Word documents
    doc_directory = r"C:\Users\lazare.kolebka\OneDrive - Accenture\Desktop\Yara"

    # List to hold the quality check results for all documents
    results = []

    # Iterate over all Word documents in the directory
    for filename in os.listdir(doc_directory):
        if filename.endswith(".docx"):
            doc_path = os.path.join(doc_directory, filename)
            result = check_document_quality(doc_path)
            results.append(result)
            print(f"Finished checking document: {doc_path}")

    # Convert results to a DataFrame
    df = pd.DataFrame(results)

    # Save results to an Excel file
    output_path = r"C:\Users\lazare.kolebka\OneDrive - Accenture\Desktop\Yara\document_quality_report.xlsx"
    try:
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            df.to_excel(writer, index=False)
        print(f"Quality check completed. Report saved to {output_path}")
    except Exception as e:
        print(f"Failed to save the report to {output_path}: {e}")

if __name__ == "__main__":
    main()
