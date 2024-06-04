import os
from docx import Document
from docx.oxml.ns import qn

def print_document_content(doc_path):
    print(f"_______________Checking content in document: {doc_path}")
    doc = Document(doc_path)

    for paragraph in doc.paragraphs:
        style = paragraph.style
        style_name = style.name if style else 'No Style'
        paragraph_text = paragraph.text.strip()
        print(f"Style: {style_name}, Text: {paragraph_text}")

    for table in doc.tables:
        print("Table detected:")
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells)
            print(row_text)

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            print("Image detected")

# Path to the directory containing the Word documents
doc_directory = r"C:\Users\lazare.kolebka\OneDrive - Accenture\Desktop\Yara"

# Iterate over all Word documents in the directory
for filename in os.listdir(doc_directory):
    if filename.endswith(".docx"):
        doc_path = os.path.join(doc_directory, filename)
        print_document_content(doc_path)
        print(f"Finished checking content in document: {doc_path}")

print("Content check completed.")
